"""Monthly report helpers for the Outreach & Admissions report."""

from __future__ import annotations

import copy
import csv
import io
from collections import Counter

from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

REPORT_TYPE_OUTREACH_ADMISSIONS = "OUTREACH_ADMISSIONS_MONTHLY_REPORT"
REPORT_TYPE_LEGACY = "LEGACY_MONTHLY_REPORT"


def month_year_label(month: int, year: int) -> str:
    month_names = [
        "January",
        "February",
        "March",
        "April",
        "May",
        "June",
        "July",
        "August",
        "September",
        "October",
        "November",
        "December",
    ]
    month_name = month_names[max(1, min(12, int(month))) - 1]
    return f"{month_name} {year}"


def _blank_totals() -> dict:
    return {
        "totalApplicationsSubmitted": 0,
        "femaleApplicants": 0,
        "maleApplicants": 0,
        "centersUtilized": 0,
        "centerDistribution": "",
        "confirmedArrivals": 0,
        "scheduledArrivals": 0,
        "applicantInterviews": 0,
        "workforceVisitsMeetings": 0,
        "schoolOutreachActivities": 0,
        "communityCivicEngagementActivities": 0,
        "campusVisits": 0,
    }


def build_blank_report_data(month: int, year: int, counselor: str = "") -> dict:
    return {
        "reportType": REPORT_TYPE_OUTREACH_ADMISSIONS,
        "title": "Outreach & Admissions Monthly Report",
        "header": {
            "oaCounselor": counselor or "",
            "monthYearReporting": month_year_label(month, year),
            "center": "",
            "countiesCovered": "",
        },
        "applicants": [
            {
                "id": "applicant-1",
                "applicant": "",
                "gender": "",
                "age": "",
                "trade": "",
                "center": "",
                "status": "",
            }
        ],
        "totals": _blank_totals(),
        "addendumGoals": {
            "obs": {"strategies": "", "outcomes": ""},
            "partnerCollaboration": {"strategies": "", "outcomes": ""},
        },
        "operationalContext": "",
        "familyEngagement": [
            {"id": "family-1", "date": "", "activity": ""}
        ],
        "outreachActivities": [
            {"id": "outreach-1", "date": "", "activity": ""}
        ],
        "communityServiceVisibility": {
            "communityServiceEvents": "",
            "marketingVisibility": "",
            "tradesBookletQrCodeUsage": "",
            "busOrBillboardAdvertising": "",
        },
        "strategicRelationshipDevelopment": [],
        "upcomingInitiatives": {
            "juneInitiatives": "",
            "julyInitiatives": "",
            "jobCorpsInformationSymposium": "",
        },
        "monthlyImpactStatement": "",
    }


def normalize_report_data(report_data: dict | None, month: int, year: int, counselor: str = "") -> dict:
    base = build_blank_report_data(month, year, counselor)
    if not report_data:
        return base

    merged = copy.deepcopy(base)
    for key, value in report_data.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            nested = dict(merged[key])
            nested.update(value)
            merged[key] = nested
        else:
            merged[key] = value

    merged["totals"] = {**_blank_totals(), **(merged.get("totals") or {})}

    applicants = merged.get("applicants") or []
    if not applicants:
        merged["applicants"] = base["applicants"]

    family_rows = merged.get("familyEngagement") or []
    if not family_rows:
        merged["familyEngagement"] = base["familyEngagement"]

    outreach_rows = merged.get("outreachActivities") or []
    if not outreach_rows:
        merged["outreachActivities"] = base["outreachActivities"]

    strategic = merged.get("strategicRelationshipDevelopment") or []
    if isinstance(strategic, str):
        strategic = [line.strip() for line in strategic.splitlines() if line.strip()]
    merged["strategicRelationshipDevelopment"] = strategic

    return merged


def derive_totals_from_applicants(applicants: list[dict]) -> dict:
    centers = []
    female_count = 0
    male_count = 0
    for row in applicants or []:
        applicant_name = (row.get("applicant") or "").strip()
        center = (row.get("center") or "").strip()
        if center:
            centers.append(center)
        gender = (row.get("gender") or "").strip().lower()
        if not applicant_name:
            continue
        if gender.startswith("f"):
            female_count += 1
        elif gender.startswith("m"):
            male_count += 1

    center_counts = Counter(centers)
    center_distribution = "\n".join(
        f"{center}: {count}" for center, count in center_counts.items()
    )

    return {
        "totalApplicationsSubmitted": len([row for row in applicants or [] if (row.get("applicant") or "").strip()]),
        "femaleApplicants": female_count,
        "maleApplicants": male_count,
        "centersUtilized": len(center_counts),
        "centerDistribution": center_distribution,
    }


def _safe_text(value) -> str:
    if value is None:
        return ""
    return str(value)


def _ensure_title(report) -> str:
    if report.report_type == REPORT_TYPE_OUTREACH_ADMISSIONS and report.report_data:
        return report.report_data.get("title") or "Outreach & Admissions Monthly Report"
    return f"Monthly Report - {month_year_label(report.month, report.year)}"


def _header_rows(report_data: dict) -> list[list[str]]:
    header = report_data.get("header", {})
    return [
        ["OA Counselor", _safe_text(header.get("oaCounselor"))],
        ["Month/Year Reporting", _safe_text(header.get("monthYearReporting"))],
        ["Center", _safe_text(header.get("center"))],
        ["Counties Covered", _safe_text(header.get("countiesCovered"))],
    ]


def _applicant_rows(report_data: dict) -> list[list[str]]:
    rows = [["Applicant", "Gender", "Age", "Trade", "Center", "Status"]]
    for applicant in report_data.get("applicants", []):
        rows.append([
            _safe_text(applicant.get("applicant")),
            _safe_text(applicant.get("gender")),
            _safe_text(applicant.get("age")),
            _safe_text(applicant.get("trade")),
            _safe_text(applicant.get("center")),
            _safe_text(applicant.get("status")),
        ])
    return rows


def _activity_rows(items: list[dict]) -> list[list[str]]:
    rows = [["Date", "Activity"]]
    for item in items or []:
        rows.append([_safe_text(item.get("date")), _safe_text(item.get("activity"))])
    return rows


def _mapping_rows(mapping: dict) -> list[list[str]]:
    rows = []
    for key, value in mapping.items():
        if isinstance(value, (dict, list)):
            continue
        rows.append([key.replace("_", " ").title(), _safe_text(value)])
    return rows


def export_structured_report_csv(report) -> str:
    data = normalize_report_data(report.report_data or {}, report.month, report.year)
    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow(["Section", "Field", "Value", "Value 2", "Value 3", "Value 4", "Value 5"])
    for field, value in _header_rows(data):
        writer.writerow(["Report Header", field, value])

    writer.writerow([])
    for row in _applicant_rows(data)[1:]:
        writer.writerow(["Applicant Breakdown", *row])

    writer.writerow([])
    for field, value in _mapping_rows(data.get("totals", {})):
        writer.writerow(["Monthly Totals", field, value])

    writer.writerow([])
    for section_name, section in (
        ("Addendum Goals - OBS", data.get("addendumGoals", {}).get("obs", {})),
        ("Addendum Goals - Partner Collaboration", data.get("addendumGoals", {}).get("partnerCollaboration", {})),
        ("Community Service & Visibility", data.get("communityServiceVisibility", {})),
        ("Upcoming Initiatives", data.get("upcomingInitiatives", {})),
    ):
        for field, value in _mapping_rows(section):
            writer.writerow([section_name, field, value])

    writer.writerow(["Operational Context", _safe_text(data.get("operationalContext", ""))])
    for activity in data.get("familyEngagement", []):
        writer.writerow(["Applicant & Family Engagement", _safe_text(activity.get("date")), _safe_text(activity.get("activity"))])
    for activity in data.get("outreachActivities", []):
        writer.writerow(["Outreach & Admissions Activities", _safe_text(activity.get("date")), _safe_text(activity.get("activity"))])
    writer.writerow(["Strategic Relationship Development", _safe_text("\n".join(data.get("strategicRelationshipDevelopment", [])))])
    writer.writerow(["Monthly Impact Statement", _safe_text(data.get("monthlyImpactStatement", ""))])
    output.seek(0)
    return output.getvalue()


def export_structured_report_excel(report) -> bytes:
    data = normalize_report_data(report.report_data or {}, report.month, report.year)
    wb = Workbook()
    ws = wb.active
    ws.title = "Report Summary"
    ws.append([_ensure_title(report)])
    ws["A1"].font = Font(bold=True, size=14)

    def add_sheet(name: str, rows: list[list[str]]):
        sheet = wb.create_sheet(title=name[:31])
        for row in rows:
            sheet.append(row)
        for col_cells in sheet.columns:
            max_length = 0
            column = col_cells[0].column_letter
            for cell in col_cells:
                max_length = max(max_length, len(_safe_text(cell.value)))
            sheet.column_dimensions[column].width = min(max_length + 4, 45)
        return sheet

    add_sheet("Header", [["Field", "Value"], * _header_rows(data)])
    add_sheet("Applicants", _applicant_rows(data))
    add_sheet("Totals", [["Field", "Value"], * _mapping_rows(data.get("totals", {}))])
    add_sheet("OBS Goals", [["Field", "Value"], * _mapping_rows(data.get("addendumGoals", {}).get("obs", {}))])
    add_sheet("Partner Goals", [["Field", "Value"], * _mapping_rows(data.get("addendumGoals", {}).get("partnerCollaboration", {}))])
    add_sheet("Operational Context", [["Operational Context"], [_safe_text(data.get("operationalContext", ""))]])
    add_sheet("Family Engagement", _activity_rows(data.get("familyEngagement", [])))
    add_sheet("Outreach Activities", _activity_rows(data.get("outreachActivities", [])))
    add_sheet("Community Visibility", [["Field", "Value"], * _mapping_rows(data.get("communityServiceVisibility", {}))])
    add_sheet("Relationships", [["Partner / Organization"], *([[partner] for partner in data.get("strategicRelationshipDevelopment", [])])])
    add_sheet("Initiatives", [["Field", "Value"], * _mapping_rows(data.get("upcomingInitiatives", {}))])
    add_sheet("Impact", [["Monthly Impact Statement"], [_safe_text(data.get("monthlyImpactStatement", ""))]])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_structured_report_docx(report) -> bytes:
    data = normalize_report_data(report.report_data or {}, report.month, report.year)
    doc = Document()
    doc.sections[0].top_margin = Inches(0.65)
    doc.sections[0].bottom_margin = Inches(0.65)
    doc.sections[0].left_margin = Inches(0.75)
    doc.sections[0].right_margin = Inches(0.75)

    doc.add_heading(_ensure_title(report), level=0)

    def add_kv_table(items: list[list[str]]):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for field, value in items:
            row = table.add_row().cells
            row[0].text = _safe_text(field)
            row[1].text = _safe_text(value)

    def add_table(headers: list[str], rows: list[list[str]]):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row_values in rows:
            row = table.add_row().cells
            for index, value in enumerate(row_values):
                row[index].text = _safe_text(value)

    doc.add_heading("Report Header", level=1)
    add_kv_table(_header_rows(data))

    doc.add_heading("Applicant Breakdown", level=1)
    add_table(
        ["Applicant", "Gender", "Age", "Trade", "Center", "Status"],
        _applicant_rows(data)[1:],
    )

    doc.add_heading("Monthly Totals", level=1)
    add_kv_table(_mapping_rows(data.get("totals", {})))

    doc.add_heading("Addendum Goals - OBS", level=1)
    add_kv_table(_mapping_rows(data.get("addendumGoals", {}).get("obs", {})))

    doc.add_heading("Addendum Goals - Partner Collaboration", level=1)
    add_kv_table(_mapping_rows(data.get("addendumGoals", {}).get("partnerCollaboration", {})))

    doc.add_heading("Operational Context", level=1)
    doc.add_paragraph(_safe_text(data.get("operationalContext", "")) or "—")

    doc.add_heading("Applicant & Family Engagement", level=1)
    add_table(["Date", "Activity"], _activity_rows(data.get("familyEngagement", []))[1:])

    doc.add_heading("Outreach & Admissions Activities", level=1)
    add_table(["Date", "Activity"], _activity_rows(data.get("outreachActivities", []))[1:])

    doc.add_heading("Community Service & Visibility", level=1)
    add_kv_table(_mapping_rows(data.get("communityServiceVisibility", {})))

    doc.add_heading("Strategic Relationship Development", level=1)
    for item in data.get("strategicRelationshipDevelopment", []):
        doc.add_paragraph(_safe_text(item), style="List Bullet")

    doc.add_heading("Upcoming Initiatives", level=1)
    add_kv_table(_mapping_rows(data.get("upcomingInitiatives", {})))

    doc.add_heading("Monthly Impact Statement", level=1)
    doc.add_paragraph(_safe_text(data.get("monthlyImpactStatement", "")) or "—")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_structured_report_pdf(report) -> bytes:
    data = normalize_report_data(report.report_data or {}, report.month, report.year)
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.65 * inch,
    )
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="StructuredSectionHeading",
            fontSize=12,
            leading=15,
            spaceAfter=6,
            textColor=colors.HexColor("#173a6a"),
            fontName="Helvetica-Bold",
        )
    )
    story = [Paragraph(_ensure_title(report), styles["Title"]), Spacer(1, 0.12 * inch)]

    def add_kv_table(title: str, rows: list[list[str]]):
        story.append(Paragraph(title, styles["StructuredSectionHeading"]))
        table = Table([["Field", "Value"], *rows], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#173a6a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.12 * inch))

    def add_table(title: str, headers: list[str], rows: list[list[str]]):
        story.append(Paragraph(title, styles["StructuredSectionHeading"]))
        table = Table([headers, *rows], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#173a6a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.12 * inch))

    add_kv_table("Report Header", _header_rows(data))
    add_table("Applicant Breakdown", ["Applicant", "Gender", "Age", "Trade", "Center", "Status"], _applicant_rows(data)[1:])
    add_kv_table("Monthly Totals", _mapping_rows(data.get("totals", {})))
    add_kv_table("Addendum Goals - OBS", _mapping_rows(data.get("addendumGoals", {}).get("obs", {})))
    add_kv_table("Addendum Goals - Partner Collaboration", _mapping_rows(data.get("addendumGoals", {}).get("partnerCollaboration", {})))

    story.append(Paragraph("Operational Context", styles["StructuredSectionHeading"]))
    story.append(Paragraph(_safe_text(data.get("operationalContext", "")) or "—", styles["BodyText"]))
    story.append(Spacer(1, 0.12 * inch))

    add_table("Applicant & Family Engagement", ["Date", "Activity"], _activity_rows(data.get("familyEngagement", []))[1:])
    add_table("Outreach & Admissions Activities", ["Date", "Activity"], _activity_rows(data.get("outreachActivities", []))[1:])
    add_kv_table("Community Service & Visibility", _mapping_rows(data.get("communityServiceVisibility", {})))

    story.append(Paragraph("Strategic Relationship Development", styles["StructuredSectionHeading"]))
    for item in data.get("strategicRelationshipDevelopment", []):
        story.append(Paragraph(_safe_text(item), styles["BodyText"]))
    story.append(Spacer(1, 0.12 * inch))

    add_kv_table("Upcoming Initiatives", _mapping_rows(data.get("upcomingInitiatives", {})))

    story.append(Paragraph("Monthly Impact Statement", styles["StructuredSectionHeading"]))
    story.append(Paragraph(_safe_text(data.get("monthlyImpactStatement", "")) or "—", styles["BodyText"]))

    document.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def export_legacy_report_csv(report) -> str:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Field", "Value"])
    writer.writerow(["Title", _ensure_title(report)])
    for field in ["summary", "applicant_activity", "outreach_activity", "communication_activity", "county_breakdown", "barriers", "performance_analysis", "next_month_strategy"]:
        writer.writerow([field.replace("_", " ").title(), getattr(report, field, "") or ""])
    output.seek(0)
    return output.getvalue()


def export_legacy_report_excel(report) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Monthly Report"
    ws.append(["Field", "Value"])
    ws.append(["Title", _ensure_title(report)])
    for field in ["summary", "applicant_activity", "outreach_activity", "communication_activity", "county_breakdown", "barriers", "performance_analysis", "next_month_strategy"]:
        ws.append([field.replace("_", " ").title(), getattr(report, field, "") or ""])
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_legacy_report_docx(report) -> bytes:
    doc = Document()
    doc.sections[0].top_margin = Inches(0.65)
    doc.sections[0].bottom_margin = Inches(0.65)
    doc.sections[0].left_margin = Inches(0.75)
    doc.sections[0].right_margin = Inches(0.75)
    doc.add_heading(_ensure_title(report), level=0)
    for field in ["summary", "applicant_activity", "outreach_activity", "communication_activity", "county_breakdown", "barriers", "performance_analysis", "next_month_strategy"]:
        doc.add_heading(field.replace("_", " ").title(), level=1)
        doc.add_paragraph(getattr(report, field, "") or "—")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_legacy_report_pdf(report) -> bytes:
    buffer = io.BytesIO()
    document = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.65 * inch, rightMargin=0.65 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SectionHeadingLegacy", fontSize=12, leading=15, spaceAfter=6, textColor=colors.HexColor("#173a6a"), fontName="Helvetica-Bold"))
    story = [Paragraph(_ensure_title(report), styles["Title"]), Spacer(1, 0.12 * inch)]
    for field in ["summary", "applicant_activity", "outreach_activity", "communication_activity", "county_breakdown", "barriers", "performance_analysis", "next_month_strategy"]:
        story.append(Paragraph(field.replace("_", " ").title(), styles["SectionHeadingLegacy"]))
        story.append(Paragraph((getattr(report, field, "") or "—").replace("\n", "<br />"), styles["BodyText"]))
        story.append(Spacer(1, 0.12 * inch))
    document.build(story)
    buffer.seek(0)
    return buffer.getvalue()